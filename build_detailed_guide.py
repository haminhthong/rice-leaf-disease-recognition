from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/HUONG_DAN_CAI_THIEN_CHI_TIET.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
BLACK = "000000"
WHITE = "FFFFFF"


def font(run, size=11, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def table_geometry(table, widths, indent=120):
    table.autofit = False
    props = table._tbl.tblPr
    tblw = props.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        props.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = props.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        props.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            props = cell._tc.get_or_add_tcPr()
            tcw = props.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                props.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            margins = props.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                props.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                item = margins.find(qn(f"w:{side}"))
                if item is None:
                    item = OxmlElement(f"w:{side}")
                    margins.append(item)
                item.set(qn("w:w"), str(value))
                item.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_props.append(OxmlElement("w:tblHeader"))
    for idx, text in enumerate(headers):
        shade(table.rows[0].cells[idx], BLUE)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), 9.2, WHITE, True)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(values):
            if ridx % 2:
                shade(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(text), 9.1)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    p.add_run(text)


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), LIGHT_BLUE)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(f"{label}: "), 10.5, DARK_BLUE, True)
    font(p.add_run(text), 10.5)
    table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    font(header.add_run("HƯỚNG DẪN CẢI THIỆN  |  RICE LEAF DISEASE RECOGNITION"), 9, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("Tài liệu kỹ thuật • 01/09/2026  |  Trang "), 9, MUTED)
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(76)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    font(p.add_run("CẨM NANG CẢI THIỆN DỰ ÁN"), 10, BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("Nhận diện bệnh trên lá lúa"), 28, DARK_BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    font(p.add_run("Lộ trình từ Portfolio Pipeline đến AI/ML có bằng chứng và sẵn sàng mở rộng"), 14, MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    font(p.add_run("Problem  →  AI/ML Correctness  →  Software Engineering  →  Production Value"), 11, BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(f"Phiên bản hướng dẫn • {date.today().strftime('%d/%m/%Y')}"), 10.5, MUTED)
    doc.add_page_break()


def content(doc):
    doc.add_heading("1. Mục tiêu và nguyên tắc", level=1)
    doc.add_paragraph(
        "Tài liệu này chuyển kết quả audit bốn tầng thành kế hoạch triển khai cụ thể. Ưu tiên quan trọng nhất "
        "không phải thêm nhiều công nghệ, mà là chứng minh mô hình giải quyết đúng bài toán bằng dữ liệu thật, "
        "split hợp lệ và metric có thể truy xuất."
    )
    callout(
        doc,
        "Kết luận hiện tại",
        "Repository đã có pipeline phần mềm tốt và 35/35 kiểm thử đạt, nhưng dữ liệu demo tổng hợp "
        "không phải bằng chứng thực tế và chưa có checkpoint/metric huấn luyện hoàn chỉnh. Không nên "
        "tuyên bố độ chính xác thực tế.",
    )
    add_table(
        doc,
        ["Tầng", "Trạng thái", "Mục tiêu cải thiện"],
        [
            ["Problem", "Đạt một phần", "Làm rõ phạm vi hai bệnh và ý nghĩa no_detection."],
            ["AI/ML correctness", "Chưa đủ bằng chứng", "Dữ liệu thật, split đủ lớn, baseline và final test."],
            ["Software Engineering", "Khá", "Tăng integration test, logging, schema và clean-clone gate."],
            ["Production/Business", "Chưa sẵn sàng", "Concurrency control, load test, security và field validation."],
        ],
        [2100, 2100, 5160],
    )

    doc.add_heading("2. Tầng Problem — Xác định đúng sản phẩm", level=1)
    doc.add_heading("2.1 Phạm vi chuẩn", level=2)
    doc.add_paragraph(
        "Mô tả khuyến nghị: Hệ thống phát hiện và định vị vùng tổn thương có đặc trưng tương tự Bạc lá lúa "
        "và Đốm nâu trên ảnh lá lúa bằng YOLOv8. Đây là công cụ hỗ trợ sàng lọc, không phải hệ thống chẩn "
        "đoán toàn diện và không thay thế chuyên gia nông nghiệp."
    )
    bullet(doc, "Không gọi kết quả là chẩn đoán bệnh nếu chưa có kiểm chứng lâm sàng/nông học.")
    bullet(doc, "Không khẳng định lá khỏe khi mô hình không phát hiện bounding box.")
    bullet(doc, "Không mở rộng kết luận sang bệnh khác, ảnh drone hoặc ảnh toàn ruộng.")
    bullet(doc, "Ghi rõ điều kiện ảnh mục tiêu: lá đủ rõ, tổn thương nhìn thấy, độ phân giải phù hợp.")

    doc.add_heading("2.2 Chuẩn hóa trạng thái đầu ra", level=2)
    add_table(
        doc,
        ["Trạng thái", "Ý nghĩa", "HTTP/API behavior"],
        [
            ["detected", "Có ít nhất một vùng đạt confidence", "200, trả danh sách detection."],
            ["no_detection", "Không có vùng thuộc hai lớp vượt ngưỡng", "200, cảnh báo không đồng nghĩa lá khỏe."],
            ["invalid_image", "Không giải mã được hoặc ảnh quá lớn", "400/413."],
            ["unsupported_input", "MIME/định dạng không hỗ trợ", "415."],
            ["model_unavailable", "Weights/config/model chưa sẵn sàng", "503."],
        ],
        [1900, 4300, 3160],
    )
    doc.add_heading("2.3 Definition of success", level=2)
    bullet(doc, "mAP50-95 và mAP50 trên validation/test.")
    bullet(doc, "Precision, recall và AP theo từng lớp; đặc biệt theo dõi false negative.")
    bullet(doc, "Latency p50/p95, kích thước model, RAM/VRAM và throughput.")
    bullet(doc, "Đánh giá trên ảnh thực địa ngoài nguồn training và review bởi người có chuyên môn.")

    doc.add_heading("3. Tầng AI/ML — Dữ liệu và chống leakage", level=1)
    doc.add_heading("3.1 Data provenance", level=2)
    number(doc, "Ghi URL, tác giả, phiên bản, ngày tải và giấy phép của từng dataset.")
    number(doc, "Giữ ZIP nguồn bất biến; lưu SHA-256 của archive và manifest.")
    number(doc, "Không commit dữ liệu nếu license không cho phép; chỉ chia sẻ sample hợp pháp.")
    number(doc, "Data Card phải ghi số ảnh, box, nguồn, class distribution, điều kiện chụp và giới hạn.")
    callout(doc, "Gate", "Không huấn luyện chính thức nếu chưa xác định được nguồn và quyền sử dụng dữ liệu.")

    doc.add_heading("3.2 Split đủ lớn", level=2)
    doc.add_paragraph(
        "Pipeline hiện chống group leakage đúng hướng, nhưng cần từ chối dataset quá nhỏ. Không chỉ kiểm tra tỷ "
        "lệ ảnh; phải kiểm tra số group độc lập và số instance mỗi lớp trong validation/test."
    )
    add_table(
        doc,
        ["Kiểm tra bắt buộc", "Ngưỡng gợi ý ban đầu", "Hành động khi không đạt"],
        [
            ["Group độc lập mỗi split", "Train ≥ 10; Val ≥ 5; Test ≥ 5", "Dừng pipeline và yêu cầu thêm dữ liệu."],
            ["Instance mỗi lớp ở val", "≥ 20", "Không cho phép model selection."],
            ["Instance mỗi lớp ở test", "≥ 20", "Không công bố metric test."],
            ["group_id qua nhiều split", "0", "Lỗi nghiêm trọng; tạo lại split."],
            ["SHA-256 qua nhiều split", "0", "Lỗi nghiêm trọng; dedup lại."],
        ],
        [2900, 2700, 3760],
    )
    doc.add_paragraph(
        "Các ngưỡng trên là gate kỹ thuật ban đầu, không phải chuẩn khoa học cố định. Với báo cáo chính thức nên "
        "xác định cỡ mẫu dựa trên độ bất định metric và phân bố thực tế."
    )

    doc.add_heading("3.3 pHash/BK-tree và annotation conflict", level=2)
    bullet(doc, "Sửa tuyên bố 'triệt tiêu leakage' thành 'giảm nguy cơ leakage'.")
    bullet(doc, "Không khẳng định BK-tree luôn O(log N); hiệu năng phụ thuộc dữ liệu và trường hợp xấu.")
    bullet(doc, "Hiệu chỉnh phash_distance bằng tập ảnh resize, JPEG, crop, đổi sáng và ảnh khác nhau nhưng giống nền.")
    bullet(doc, "Đo precision/recall của near-duplicate matching ở ngưỡng 0, 2, 4, 6, 8.")
    bullet(doc, "Ảnh trùng SHA nhưng annotation xung đột phải được quarantine, không tự chọn nhãn có nhiều box hơn.")

    doc.add_heading("4. Tầng AI/ML — Thí nghiệm và metric", level=1)
    doc.add_heading("4.1 Baseline và candidate", level=2)
    bullet(doc, "Đổi yolov8s_champion.yaml thành yolov8s_candidate.yaml cho đến khi có kết quả.")
    bullet(doc, "Baseline chính: YOLOv8n; candidate: YOLOv8s trên cùng split, seed, image size và epoch budget.")
    bullet(doc, "Không tuning bằng test. Chỉ dùng validation để chọn model và confidence threshold.")
    bullet(doc, "Test chỉ chạy sau khi khóa config và model artifact.")

    doc.add_heading("4.2 Quy trình thực nghiệm", level=2)
    for text in (
        "Audit dữ liệu và khóa manifest checksum.",
        "Huấn luyện YOLOv8n baseline.",
        "Đánh giá validation, lưu metric và error cases.",
        "Huấn luyện YOLOv8s candidate trên cùng protocol.",
        "Chọn bằng mAP50-95 kết hợp recall từng lớp và latency.",
        "Khóa model/config/commit; đánh giá test đúng một lần.",
        "Cập nhật Model Card và bảng kết quả có nguồn artifact.",
    ):
        number(doc, text)

    doc.add_heading("4.3 Metadata cần lưu cho mỗi run", level=2)
    add_table(
        doc,
        ["Nhóm", "Trường cần lưu"],
        [
            ["Source", "git_commit, config_path/config_sha256, code version"],
            ["Data", "manifest_sha256, archive checksums, split counts"],
            ["Environment", "Python, PyTorch, Ultralytics, CUDA, GPU"],
            ["Training", "seed, model, epochs, batch, imgsz, optimizer, lr, duration"],
            ["Result", "best.pt checksum, metrics.json, per_class_metrics.csv, plots"],
        ],
        [2100, 7260],
    )

    doc.add_heading("4.4 Metric và phân tích lỗi", level=2)
    bullet(doc, "Primary metric: mAP50-95; secondary: mAP50, precision, recall, AP theo lớp.")
    bullet(doc, "Bổ sung F1 theo confidence, PR curve, confusion matrix và false negatives theo lớp.")
    bullet(doc, "Báo cáo metric theo nguồn dataset để phát hiện model học đặc trưng nguồn.")
    bullet(doc, "Không chọn model chỉ vì mAP trung bình cao nếu recall một bệnh quá thấp.")
    bullet(doc, "Không đưa số liệu vào CV nếu không truy ngược được về metrics.json, config và commit.")

    doc.add_heading("4.5 Đồng bộ training và inference", level=2)
    doc.add_paragraph(
        "Thêm inference.image_size vào config và truyền imgsz cho RiceLeafDetector. API, CLI và Streamlit phải dùng "
        "cùng image size, confidence và IoU từ một cấu hình. Ultralytics phụ trách letterbox/normalization; không "
        "tạo preprocessing riêng nếu không cần."
    )

    doc.add_heading("5. Tầng Software Engineering", level=1)
    doc.add_heading("5.1 Tách trách nhiệm API vừa đủ", level=2)
    add_table(
        doc,
        ["Tệp đề xuất", "Trách nhiệm"],
        [
            ["app/settings.py", "ApiSettings và kiểm tra biến môi trường."],
            ["app/dependencies.py", "Khởi tạo/cache detector."],
            ["app/schemas.py", "Pydantic request/response schema."],
            ["app/validation.py", "Giới hạn file, decode, pixel/magic-byte validation."],
            ["app/api.py", "Khai báo route và orchestration ngắn gọn."],
        ],
        [2600, 6760],
    )
    callout(doc, "Nguyên tắc", "Không thêm repository/service pattern nếu không có database hoặc business layer thật.")

    doc.add_heading("5.2 API schema và lỗi", level=2)
    bullet(doc, "Dùng Pydantic response model thay dict tự do.")
    bullet(doc, "Validate confidence/IoU và model path khi startup; env sai không được gây lỗi 500 khó hiểu.")
    bullet(doc, "Tách liveness và readiness; readiness phải kiểm tra model có nạp được.")
    bullet(doc, "Không trả traceback hoặc đường dẫn máy chủ cho client.")

    doc.add_heading("5.3 Logging", level=2)
    bullet(doc, "Dùng logging thay print ở prepare/train/evaluate/API.")
    bullet(doc, "Log run_id, model version, latency, status code và lỗi đã chuẩn hóa.")
    bullet(doc, "Không log nội dung ảnh, secret hoặc request body.")
    bullet(doc, "Production nên dùng JSON formatter và request ID.")

    doc.add_heading("5.4 Test cần bổ sung", level=2)
    add_table(
        doc,
        ["Nhóm", "Test bắt buộc"],
        [
            ["Data", "Dataset ít group; thiếu lớp; group/SHA cross-split; image-label mismatch."],
            ["Dedup", "Exact duplicate; near duplicate; false merge; annotation conflict quarantine."],
            ["API", "File >10 MB; MIME giả; ảnh hỏng; model 503; no_detection; success mock."],
            ["Inference", "Box rỗng; class lạ; model result mock; imgsz/config đồng bộ."],
            ["Integration", "Checkpoint nhỏ; CLI predict; Docker health/readiness."],
        ],
        [2200, 7160],
    )

    doc.add_heading("5.5 CI và clean clone", level=2)
    bullet(doc, "CI phải chạy ruff format --check src app scripts tests.")
    bullet(doc, "CI phải chạy ruff check, pytest, pip check và import smoke test.")
    bullet(doc, "Thêm Docker build + health smoke test ở job riêng.")
    bullet(doc, "Clean clone phải cài được bằng pip install -e \".[app,dev]\" trong Python 3.11 mới.")
    bullet(doc, "Không commit egg-info, pytest temp, cache, data processed, runs, artifacts hoặc weights.")

    doc.add_heading("6. Tầng Production và Business Value", level=1)
    doc.add_heading("6.1 Concurrency", level=2)
    doc.add_paragraph(
        "YOLO inference là tác vụ đồng bộ và nặng. Không chạy trực tiếp trong async event loop. Dùng threadpool cho "
        "bản đơn giản và semaphore để giới hạn số inference đồng thời theo GPU/CPU. Khi queue đầy, trả 429/503 "
        "cùng Retry-After thay vì nhận vô hạn request."
    )
    add_table(
        doc,
        ["Mức tải", "Thiết kế phù hợp"],
        [
            ["Demo cá nhân", "Một FastAPI process, model cache, threadpool, concurrency=1."],
            ["Nhóm nhỏ", "Nhiều API worker nhưng inference queue dùng chung; benchmark CPU/GPU."],
            ["100 users", "API gateway, rate limit, queue, inference worker/replica và monitoring."],
            ["Quy mô lớn", "Model serving chuyên dụng, autoscaling và artifact registry."],
        ],
        [2300, 7060],
    )

    doc.add_heading("6.2 Load test", level=2)
    bullet(doc, "Chạy 1, 10, 25, 50 và 100 user đồng thời bằng Locust/k6.")
    bullet(doc, "Đo throughput, p50/p95/p99, error rate, queue wait, CPU/RAM/GPU/VRAM.")
    bullet(doc, "Gate gợi ý: error <1%, không OOM/crash, overload trả 429 rõ ràng.")
    bullet(doc, "Không ghi 'hỗ trợ 100 users' trong README nếu chưa có report load test.")

    doc.add_heading("6.3 Security và privacy", level=2)
    bullet(doc, "Không tin MIME từ client; kiểm tra magic bytes và decode thực tế.")
    bullet(doc, "Giới hạn tổng pixel để chống decompression bomb.")
    bullet(doc, "Bổ sung rate limiting, timeout, CORS allowlist và authentication nếu API public.")
    bullet(doc, "Chạy dependency audit; đánh giá vulnerability thay vì tự động ignore.")
    bullet(doc, "Không lưu ảnh mặc định; nếu lưu phải có consent, retention và xóa EXIF khi phù hợp.")
    bullet(doc, "Manifest công khai chỉ dùng đường dẫn tương đối/ID, không chứa path máy cá nhân.")

    doc.add_heading("6.4 Chứng minh business value", level=2)
    bullet(doc, "Demo bằng ảnh thật ngoài training, có cả success, no_detection và ảnh khó.")
    bullet(doc, "Nhờ chuyên gia nông nghiệp review false positive/false negative.")
    bullet(doc, "Đo thời gian xử lý thủ công so với hệ thống và chi phí inference.")
    bullet(doc, "Xác định quyết định mà hệ thống hỗ trợ; tránh chỉ trình bày bounding box đẹp.")

    doc.add_heading("7. README, Model Card và CV", level=1)
    bullet(doc, "README ưu tiên bài toán, data, kiến trúc, cài đặt, chạy, kết quả, hạn chế và reproducibility.")
    bullet(doc, "Chuyển câu hỏi phỏng vấn sang docs/INTERVIEW_NOTES.md để README gọn hơn.")
    bullet(doc, "Gọi ảnh tổng hợp là synthetic smoke-test asset, không phải bằng chứng accuracy.")
    bullet(doc, "Model Card phải có metric từng lớp, dataset version, intended use và failure modes.")
    bullet(doc, "Chỉ ghi số mAP/latency trong CV khi có artifact chứng minh.")
    add_table(
        doc,
        ["Model", "Val mAP50-95", "Test mAP50-95", "Recall BLB", "Recall Brown Spot", "p95 latency"],
        [["YOLOv8n baseline", "TBD", "TBD", "TBD", "TBD", "TBD"], ["YOLOv8s candidate", "TBD", "TBD", "TBD", "TBD", "TBD"]],
        [1800, 1450, 1450, 1450, 1760, 1450],
    )

    doc.add_heading("8. Lộ trình triển khai", level=1)
    add_table(
        doc,
        ["Phase", "Công việc", "Điều kiện hoàn thành"],
        [
            ["1. Truthfulness", "Sửa claim BK-tree/leakage; rename champion; làm rõ no_detection; CI lint scripts.", "Không còn claim chưa có bằng chứng; test/lint xanh."],
            ["2. Data correctness", "Provenance/license; split gates; conflict quarantine; pHash calibration; manifest hash.", "Val/test đủ lớp; 0 group/SHA leakage."],
            ["3. Real experiments", "Train baseline/candidate; val selection; final test; error analysis.", "Metric tái lập và truy được về artifact."],
            ["4. Reliability", "Schema, config đồng bộ, logging, edge/integration tests, readiness.", "Clean clone và Docker smoke test đạt."],
            ["5. Production evidence", "Concurrency, rate limit, security, load test, monitoring.", "Có report tải và không tuyên bố quá khả năng."],
        ],
        [1500, 4460, 3400],
    )

    doc.add_heading("9. Definition of Done trước khi đưa vào CV", level=1)
    checklist = [
        "Bài toán và phạm vi hai bệnh được mô tả chính xác.",
        "Dataset có nguồn, license và checksum.",
        "Không phát hiện exact/group leakage giữa split.",
        "Validation/test đủ group và instance mỗi lớp.",
        "Có YOLOv8n baseline thật và candidate thật.",
        "Có validation metric, final test metric và error analysis.",
        "Metric truy được về config, manifest, commit và checkpoint.",
        "Training/inference dùng cùng image size và preprocessing.",
        "Ảnh demo thực không xuất hiện trong training.",
        "API xử lý input sai, model thiếu và overload.",
        "35 test hiện tại vẫn đạt và test mới bao phủ thêm các trường hợp biên/tích hợp.",
        "Clean clone, CI và Docker smoke test đạt.",
        "Không có secret, path cá nhân hoặc artifact tạm trong Git.",
        "README và Model Card không có claim quá mức.",
        "Nếu nói hỗ trợ nhiều người dùng, có load-test report.",
    ]
    for item in checklist:
        bullet(doc, "☐ " + item)

    doc.add_heading("10. Ưu tiên hành động ngay", level=1)
    callout(
        doc,
        "P0",
        "Sửa tính trung thực của tài liệu: BK-tree không bảo đảm O(log N), pHash chỉ giảm nguy cơ leakage, "
        "candidate chưa phải champion, no_detection không phải lá khỏe.",
    )
    callout(
        doc,
        "P1",
        "Chuẩn bị dataset thật có license; thêm split quality gates; huấn luyện YOLOv8n baseline và YOLOv8s candidate.",
    )
    callout(
        doc,
        "P2",
        "Bổ sung integration test, readiness, logging và đồng bộ imgsz trước khi tối ưu production.",
    )
    callout(
        doc,
        "P3",
        "Chỉ triển khai queue/autoscaling sau khi load test chứng minh nhu cầu và model đã có giá trị thực nghiệm.",
    )


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    cover(doc)
    content(doc)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
